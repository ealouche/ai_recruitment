"""
Module d'extraction de texte optimisé pour les CVs
Auteur: Elyes Alouache
Date: 09/08/2025

Optimisé pour un usage minimal des ressources (CPU, mémoire, E/S).
Extraction rapide et linéaire avec préservation du formatage.
"""

import io
import logging
from typing import Optional, Union
from pathlib import Path

# Imports conditionnels pour optimiser les performances
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

logger = logging.getLogger(__name__)


class OptimizedTextExtractor:
    """
    Extracteur de texte optimisé pour les CVs.
    
    Caractéristiques d'optimisation :
    - Lecture en streaming pour minimiser l'usage mémoire
    - Extraction linéaire sans parsing inutile
    - Préservation des sauts de ligne pour la lisibilité
    - Gestion d'erreurs robuste
    - Support multi-format (PDF, DOCX, TXT)
    """
    
    @staticmethod
    def extract_text_from_bytes(file_content: bytes, filename: str) -> str:
        """
        Extrait le texte d'un fichier à partir de son contenu en bytes.
        
        Args:
            file_content: Contenu du fichier en bytes
            filename: Nom du fichier pour déterminer le type
            
        Returns:
            Texte extrait avec sauts de ligne préservés
            
        Raises:
            ValueError: Si le format n'est pas supporté
            Exception: Pour les erreurs d'extraction
        """
        if not file_content:
            raise ValueError("Contenu du fichier vide")
        
        # Détection du type de fichier basée sur l'extension
        file_extension = Path(filename).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return OptimizedTextExtractor._extract_from_pdf(file_content)
            elif file_extension in ['.docx', '.doc']:
                return OptimizedTextExtractor._extract_from_docx(file_content)
            elif file_extension == '.txt':
                return OptimizedTextExtractor._extract_from_txt(file_content)
            else:
                # Tentative d'extraction PDF par défaut (beaucoup de CVs sont en PDF)
                try:
                    return OptimizedTextExtractor._extract_from_pdf(file_content)
                except:
                    raise ValueError(f"Format de fichier non supporté: {file_extension}")
                    
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction de texte de {filename}: {str(e)}")
            raise Exception(f"Impossible d'extraire le texte: {str(e)}")
    
    @staticmethod
    def _extract_from_pdf(file_content: bytes) -> str:
        """
        Extraction optimisée de texte PDF.
        Utilise pdfplumber en priorité pour une meilleure qualité,
        fallback sur PyPDF2 pour la compatibilité.
        """
        if not PDF_AVAILABLE and not PYPDF2_AVAILABLE:
            raise ValueError("Aucune bibliothèque PDF disponible")
        
        # Méthode 1: pdfplumber (meilleure qualité, préserve le formatage)
        if PDF_AVAILABLE:
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            # Préservation des sauts de ligne et nettoyage minimal
                            cleaned_text = page_text.strip()
                            if cleaned_text:
                                text_parts.append(cleaned_text)
                    
                    return '\n\n'.join(text_parts) if text_parts else ""
            except Exception as e:
                logger.warning(f"pdfplumber a échoué, tentative avec PyPDF2: {str(e)}")
        
        # Méthode 2: PyPDF2 (fallback)
        if PYPDF2_AVAILABLE:
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                text_parts = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        cleaned_text = page_text.strip()
                        if cleaned_text:
                            text_parts.append(cleaned_text)
                
                return '\n\n'.join(text_parts) if text_parts else ""
            except Exception as e:
                raise Exception(f"Erreur PyPDF2: {str(e)}")
        
        raise ValueError("Aucune méthode d'extraction PDF disponible")
    
    @staticmethod
    def _extract_from_docx(file_content: bytes) -> str:
        """
        Extraction optimisée de texte DOCX.
        Lecture linéaire des paragraphes avec préservation du formatage.
        """
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx non disponible")
        
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []
            
            # Extraction linéaire des paragraphes
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extraction des tableaux (souvent utilisés dans les CVs)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n\n'.join(text_parts) if text_parts else ""
            
        except Exception as e:
            raise Exception(f"Erreur extraction DOCX: {str(e)}")
    
    @staticmethod
    def _extract_from_txt(file_content: bytes) -> str:
        """
        Extraction optimisée de texte TXT.
        Gestion de l'encodage avec fallback.
        """
        try:
            # Tentative UTF-8 en premier
            return file_content.decode('utf-8').strip()
        except UnicodeDecodeError:
            try:
                # Fallback sur latin-1
                return file_content.decode('latin-1').strip()
            except UnicodeDecodeError:
                # Fallback sur cp1252 (Windows)
                try:
                    return file_content.decode('cp1252').strip()
                except UnicodeDecodeError:
                    # Dernier recours: ignorer les caractères problématiques
                    return file_content.decode('utf-8', errors='ignore').strip()
    
    @staticmethod
    def validate_extracted_text(text: str) -> bool:
        """
        Valide que le texte extrait est utilisable.
        Utilise les paramètres de configuration optimisés.
        
        Args:
            text: Texte à valider
            
        Returns:
            True si le texte est valide, False sinon
        """
        from app.core.config import settings
        
        if not text or not text.strip():
            return False
        
        # Vérification de la longueur minimale (configurable)
        if len(text.strip()) < settings.MIN_TEXT_LENGTH:
            return False
        
        # Vérification qu'il y a des mots (configurable)
        words = text.split()
        if len(words) < settings.MIN_WORD_COUNT:
            return False
        
        return True
    
    @staticmethod
    def get_text_stats(text: str) -> dict:
        """
        Retourne des statistiques sur le texte extrait.
        
        Args:
            text: Texte à analyser
            
        Returns:
            Dictionnaire avec les statistiques
        """
        if not text:
            return {
                "character_count": 0,
                "word_count": 0,
                "line_count": 0,
                "is_valid": False
            }
        
        lines = text.split('\n')
        words = text.split()
        
        return {
            "character_count": len(text),
            "word_count": len(words),
            "line_count": len(lines),
            "non_empty_lines": len([line for line in lines if line.strip()]),
            "is_valid": OptimizedTextExtractor.validate_extracted_text(text)
        }
